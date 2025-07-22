from typing import Any
import httpx
from mcp.server.fastmcp import FastMCP
import asyncio
import shutil
import os
import re
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber
from pathlib import Path

# Initialize FastMCP server
mcp = FastMCP("weather")

# Constants
NWS_API_BASE = "https://api.weather.gov"
USER_AGENT = "weather-app/1.0"

@mcp.tool()
async def get_weather_data(location: str) -> Any:
    """
    Fetches weather data for a given location using the NWS API.
    
    Args:
        location (str): The location for which to fetch weather data.
        
    Returns:
        Any: The weather data in JSON format.
    """
    headers = {"User-Agent": USER_AGENT}
    async with httpx.AsyncClient() as client:
        response = await client.get(f"{NWS_API_BASE}/points/{location}", headers=headers)
        response.raise_for_status()
        return response.json()

@mcp.tool()
async def duplicate_folder(source_folder: str, new_folder: str) -> Any:
    """
    Duplicates a folder with a new folder name.

    Args:
        source_folder (str): The path to the folder to duplicate.
        new_folder (str): The path for the new duplicated folder.

    Returns:
        Any: A message indicating success or error.
    """
    try:
        if not os.path.exists(source_folder):
            return {"error": f"Source folder '{source_folder}' does not exist."}
        if os.path.exists(new_folder):
            return {"error": f"Destination folder '{new_folder}' already exists."}
        shutil.copytree(source_folder, new_folder)
        return {"message": f"Folder duplicated from '{source_folder}' to '{new_folder}'."}
    except Exception as e:
        return {"error": str(e)}

@mcp.tool()
async def change_month_in_name(folder_path: str) -> Any:
    """
    Renames all files in a folder by replacing the month in their names with the current month (uppercase, Spanish).

    Args:
        folder_path (str): The path to the folder containing the files.

    Returns:
        Any: A list of messages indicating success or error for each file.
    """
    try:
        if not os.path.isdir(folder_path):
            return {"error": f"Folder '{folder_path}' does not exist."}

        # List of month names in English and Spanish
        months = [
            "january", "february", "march", "april", "may", "june",
            "july", "august", "september", "october", "november", "december",
            "enero", "febrero", "marzo", "abril", "mayo", "junio",
            "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre"
        ]
        # Spanish months in order
        spanish_months = [
            "ENERO", "FEBRERO", "MARZO", "ABRIL", "MAYO", "JUNIO",
            "JULIO", "AGOSTO", "SEPTIEMBRE", "OCTUBRE", "NOVIEMBRE", "DICIEMBRE"
        ]
        # Get current month index (1-12)
        current_month_index = datetime.now().month - 1
        current_month_spanish = spanish_months[current_month_index]

        pattern = re.compile(r'(' + '|'.join(months) + r')', re.IGNORECASE)

        results = []
        for filename in os.listdir(folder_path):
            file_path = os.path.join(folder_path, filename)
            if os.path.isfile(file_path):
                if pattern.search(filename):
                    new_filename = pattern.sub(current_month_spanish, filename, count=1)
                    new_file_path = os.path.join(folder_path, new_filename)
                    try:
                        os.rename(file_path, new_file_path)
                        results.append({"message": f"Renamed '{filename}' to '{new_filename}'."})
                    except Exception as e:
                        results.append({"error": f"Error renaming '{filename}': {str(e)}"})
        return results
    except Exception as e:
        return {"error": str(e)}

import os
from typing import Any
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.text.paragraph import Paragraph
from docx.shared import Inches

from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_borders(table):
    tbl = table._element

    # Obtener o crear <w:tblPr>
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    # Crear <w:tblBorders>
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')     # Línea simple
        border.set(qn('w:sz'), '5')           # Grosor (1pt)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '#bcbcbc')   # Negro
        tblBorders.append(border)

    # Añadir bordes a tblPr (eliminando los anteriores si existen)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def set_cell_background(cell, color_hex):
    """
    Establece el color de fondo de una celda en formato hexadecimal (ej. '0070C0').
    """
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def create_table_with_header(doc, headers, widths):
            table = doc.add_table(rows=1, cols=len(headers))
            tbl_pr = table._element.xpath('./w:tblPr')[0]
            jc = OxmlElement('w:jc')
            jc.set(qn('w:val'), 'center')
            tbl_pr.append(jc)
            if 'Table Grid' in [s.name for s in doc.styles]:
                table.style = 'Table Grid'
            else:
                set_borders(table)
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                para = hdr_cells[i].paragraphs[0]
                run = para.add_run(header)
                run.bold = True
                hdr_cells[i].width = widths[i]
                set_cell_background(hdr_cells[i], "#f4f4f4")
            return table

def insert_title_table(doc, parent, index):
            p_title = OxmlElement('w:p')
            r_title = OxmlElement('w:r')
            t_title = OxmlElement('w:t')
            r_title.append(t_title)
            p_title.append(r_title)
            parent.insert(index, p_title)
            return Paragraph(p_title, doc)

@mcp.tool()
async def generate_monthly_report(
    report_filepath: str,
    month: str,
    activities: str,
    conclusiones: str,
    recommendations: str,
    title_activities: list,  # Lista de diccionarios con "actividad" y "mes"
    description_activities: list  # <--- Nuevo parámetro agregado
) -> Any:
    """
    Abre un archivo DOC o DOCX, reemplaza los marcadores con texto o tabla.
    Reemplaza {activities}, {month}, {conclusions}, {recommendations} con texto justificado.
    Reemplaza {titleActivities} con una tabla generada desde title_activities.
    Reemplaza {descriptionActivities} con una tabla generada desde description_activities.

    Args:
        report_filepath (str): Ruta del archivo .doc o .docx.
        month (str): Mes a insertar.
        activities (str): Texto de actividades.
        conclusiones (str): Texto de conclusiones.
        recommendations (str): Texto de recomendaciones.
        title_activities (list): Lista de dicts con "actividad" y "mes".
        description_activities (list): Lista de dicts con "actividad", "descripcion" y "verifcador".

    Returns:
        Any: Mensaje de éxito o error.
    """
    try:
        if not os.path.isfile(report_filepath):
            return {"error": f"File '{report_filepath}' does not exist."}

        ext = os.path.splitext(report_filepath)[1].lower()
        if ext not in [".doc", ".docx"]:
            return {"error": "Only DOC or DOCX files are supported."}

        doc = Document(report_filepath)
        replaced = False
        
        for i, para in enumerate(doc.paragraphs):
            if "{titleActivities}" in para.text or "{descriptionActivities}" in para.text:
                parent = para._element.getparent()
                index = parent.index(para._element)
                parent.remove(para._element)

                if "{titleActivities}" in para.text:
                    insert_title_table(doc, parent, index)
                    col1_width, col2_width = Inches(4.5), Inches(1.5)
                    table = create_table_with_header(doc, ["ACTIVIDADES", "MESES"], [col1_width, col2_width])
                    for item in title_activities:
                        row = table.add_row().cells
                        row[0].text = item["actividad"]
                        row[1].text = item["mes"]
                        row[0].width = col1_width
                        row[1].width = col2_width
                    parent.insert(index + 1, table._element)

                elif "{descriptionActivities}" in para.text:
                    insert_title_table(doc, parent, index)
                    col_widths = [Inches(2.5), Inches(3.0), Inches(1.5)]
                    headers = ["Actividad Planificada", "Actividad Ejecutada", "Verificador"]
                    table = create_table_with_header(doc, headers, col_widths)
                    for item in description_activities:
                        row = table.add_row().cells
                        row[0].text = item["actividad"]
                        row[1].text = item["descripcion"]
                        row[2].text = item["verificador"]
                        for j, width in enumerate(col_widths):
                            row[j].width = width
                    parent.insert(index + 1, table._element)
                replaced = True

        # Reemplazo en párrafos
        for para in doc.paragraphs:
            if any(x in para.text for x in ["{activities}", "{month}", "{conclusions}", "{recommendations}"]):
                if "{activities}" in para.text:
                    para.text = para.text.replace("{activities}", activities)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if "{month}" in para.text:
                    para.text = para.text.replace("{month}", month)
                if "{conclusions}" in para.text:
                    para.text = para.text.replace("{conclusions}", conclusiones)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if "{recommendations}" in para.text:
                    para.text = para.text.replace("{recommendations}", recommendations)
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                replaced = True

        # Reemplazo en tablas existentes
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if any(x in cell.text for x in ["{activities}", "{month}", "{conclusions}", "{recommendations}"]):
                        if "{activities}" in cell.text:
                            cell.text = cell.text.replace("{activities}", activities)
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        if "{month}" in cell.text:
                            cell.text = cell.text.replace("{month}", month)
                        if "{conclusions}" in cell.text:
                            cell.text = cell.text.replace("{conclusions}", conclusiones)
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        if "{recommendations}" in cell.text:
                            cell.text = cell.text.replace("{recommendations}", recommendations)
                            for p in cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        replaced = True

        if replaced:
            doc.save(report_filepath)
            return {"message": f"Replaced placeholders and saved '{report_filepath}'."}
        else:
            return {"message": "No placeholders found in the document."}
    except Exception as e:
        return {"error": str(e)}


@mcp.tool()
async def extract_text_from_pdfs(folder_path: str) -> Any:
    """
    Extracts all text from PDF files in the given folder using pdf-plumber.

    Args:
        folder_path (str): The path to the folder containing PDF files.

    Returns:
        Any: A dictionary with filenames as keys and extracted text as values, or an error message.
    """
    try:
        if not os.path.isdir(folder_path):
            return {"error": f"Folder '{folder_path}' does not exist."}

        pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith('.pdf')]
        if not pdf_files:
            return {"error": "No PDF files found in the folder."}

        extracted = {}
        for filename in pdf_files:
            file_path = os.path.join(folder_path, filename)
            try:
                with pdfplumber.open(file_path) as pdf:
                    text = ""
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                    extracted[filename] = text
            except Exception as e:
                extracted[filename] = f"Error extracting text: {str(e)}"
        return extracted
    except Exception as e:
        return {"error": str(e)}
    

@mcp.tool()
def generate_monthly_report_prompt(report_path: str, folder_path: str) -> str:
    """
    Genera un prompt estructurado para que un agente de IA lea documentos PDF desde una carpeta
    y genere un informe mensual con base en dichos documentos.

    Parámetros:
    - report_path: Ruta del archivo donde se generará el informe mensual.
    - folder_path: Carpeta donde se encuentran los archivos PDF a resumir.

    Retorna:
    - Un prompt con instrucciones detalladas para el agente.
    """

    report_file = Path(report_path).name
    folder_name = Path(folder_path).name

    prompt = f"""
***INSTRUCCIONES:***

NO redactes, resumas ni interpretes este mensaje.  
NO añadas contexto, comentarios ni justificaciones.  
Tu única tarea es ejecutar las acciones descritas en este texto exactamente como están escritas.  

A continuación, se te proporciona el prompt a seguir de forma literal:

---
Sigue los siguientes pasos cuidadosamente para generar el reporte mensual:

1. En la carpeta "{folder_path}", lee todos los archivos en formato PDF. Por cada archivo, genera un resumen con el siguiente formato:

   > Mediante {{tipo de documento: Memorando, Oficio, Correo}}, Nro. {{número de documento}}, de fecha {{fecha del documento}}, se indica que {{resumen del contenido}}.

2. Usa los resúmenes generados en el paso anterior para completar el archivo "{report_path}", correspondiente al reporte mensual. En este archivo debes:
   - Incluir las actividades descritas en los resúmenes.
   - A partir del contenido de los resúmenes generados, construye una lista llamada `title_activities`. Esta lista debe estar en formato dictionary de python y contener objetos, uno por cada actividad mencionada en los resúmenes con las siguientes claves:
     - `"actividad"`: un asunto de la actividad mencionada en el documento.
     - `"mes"`: el mes en el que ocurre o se menciona dicha actividad.

        Ejemplo de estructura esperada:

        ```python
        title_activities = [
            {{"actividad": "Entrega de informe técnico a la Subsecretaría", "mes": "Marzo"}},
            {{"actividad": "Capacitación sobre Registro Social", "mes": "Marzo"}}
        ]
        ```
   - A partir del contenido de los resúmenes generados, construye una lista llamada `description_activities`. Esta lista debe estar en formato dictionary de python y contener objetos, uno por cada actividad mencionada en los resúmenes con las siguientes claves:
     - `"actividad"`: es la misma actividad generadad en la variable `title_activities`.
     - `"descripcion"`: el un resumen corto de la actividad realizada obtenido de los resúmenes.
     - `"verificador"`: es el numero y tipo de documento obtenido del resumen, sea el oficio, memorando o correo electronico.

        Ejemplo de estructura esperada:

        ```python
        title_activities = [
            {{"actividad": "Entrega de informe técnico a la Subsecretaría", "descripcion":"Se convoca a la Fundación Naveducando a una reunión el 23 de enero de 2025 a las 15h00 (hora Galápagos), para revisar las actividades planificadas del convenio CGREG-CONV-010-2021, de forma presencial o vía Zoom.", "verificador": "Oficio Nro. CGREG-DDFPES-2025-0014-OF"}},
            {{"actividad": "Capacitación sobre Registro Social", "descripcion": "Se convoca a la Comisión Técnica y a ESPOLTECH-EP a una reunión del proceso RE-CEP-CGREG-2025-001 el 16 de junio de 2025 a las 15h00 (hora insular) por Zoom, para tratar la designación del secretario/a y preguntas del proceso.", "verificador": "Memorando Nro. CGREG-DDFPES-2025-0319-M"}},
            {{"actividad": "Capacitación sobre Registro Social", "descripcion": "Se informa que el 25 de julio a las 17h30 en San Cristóbal se realizará un evento de inicio de la consultoría liderada por la DDFPES, solicitando apoyo con infocus y pantalla para el taller.", "verificador": "Correo electrónico"}}
        ]
        ```
   - Agregar una conclusión que contenga dos párrafos:
     - El primer párrafo debe ser breve (entre 3 a 5 líneas) y debe sintetizar las actividades ejecutadas durante el mes.
     - El segundo párrafo debe tener la siguiente estructura:  
       "Durante el mes de {{mes}} del {{año}}, se generaron los siguientes productos: {{lista de productos}}."  
       En esta lista deben incluirse productos mencionados en las actividades, tales como informes, ayudas memoria, convenios suscritos, etc.  
       Si no se menciona ningún producto en las actividades, entonces escribe: "no se generaron productos".
   - Proponer una recomendación basada en las actividades realizadas.
   - El **mes y año** del reporte deben ser obtenidos del nombre del archivo "{report_file}". Utiliza esta información para titular y contextualizar adecuadamente el reporte.
Asegúrate de que los resúmenes sean claros, concisos y mantengan coherencia con el contenido original. El prompt debe ser identico al texto proporcionado, sin cambios ni adiciones.
"""

    return prompt.strip()


if __name__ == "__main__":
    asyncio.run(mcp.run())
